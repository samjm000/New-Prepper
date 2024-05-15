import docx
import doc2docx
import os
import excel_processor
import document_organiser
import document_converter
from collections import OrderedDict
from ClinicClass import ClinicPatient
from datetime import datetime


# Convert date from DDMMYY to datetime object
def convert_date(date_string):
    try:
        return datetime.strptime(date_string, "%d%m%y")
    except ValueError:
        return datetime.min  # Return None if date_string does not match format


# Find summary from most recent document for each CHI
def patient_details_finder(clinic_list, document_list):
    document_list.sort(key=lambda doc: convert_date(doc[-6:]), reverse=True)
    document_list = document_converter.process_documents(document_list)
    list_of_patients = []

    # create a list of patients
    for CHI, clinic_info in clinic_list.items():
        #extract clinic details from clnic_info
        appt_time = clinic_info['time']
        f_name = clinic_info['fname']
        l_name = clinic_info['lname']
        clinic_patient = ClinicPatient(CHI, appt_time, f_name, l_name)
        list_of_patients.append(clinic_patient)

    # for every patient look for CHI
    for clinic_patient in list_of_patients:
        search_CHI = clinic_patient.CHI
        print(f"Looking for {search_CHI}")
        found = False  # Flag to check if CHI found

        # For each document
        for document in document_list:
            if found:
                break
            current_document = docx.Document(document)
            # For each CHI - find details
            for i, paragraph in enumerate(current_document.paragraphs):
                if search_CHI in paragraph.text:
                    # Save text
                    current_paragraph = paragraph.text
                    next_paragraph_index = i + 1
                    # Read through rest of document until ':' found
                    while next_paragraph_index < len(current_document.paragraphs):
                        next_paragraph = current_document.paragraphs[
                            next_paragraph_index
                        ].text
                        if ":" in next_paragraph:
                            break
                        current_paragraph += "\n" + next_paragraph
                        next_paragraph_index += 1

                    # print(f"Adding this string to patient summary {current_paragraph}")
                    clinic_patient.summary = current_paragraph
                    clinic_patient.last_seen = document[-11:-5]
                    found = True

    return list_of_patients


if __name__ == "__main__":
    
    print("Running Test")

    # Load test data
    # Importing Excel file details
    input_excel_file = r"clinic lists\GN Clinic 200524.xls"
    print(f"Reading in File : {input_excel_file}")
    clinic_list = excel_processor.get_chi_numbers(input_excel_file)

    print(f"Generated time and CHI dictionary {clinic_list}\\n")

    # organise and get list of previous documents (Not dealt with .doc files yet)
    sorted_document_list = document_organiser.get_documents_in_order()

    # get new document list converting .doc to .docx
    all_documents_docx = document_converter.process_documents(sorted_document_list)

    # import patient clinic list details
    patient_details_finder(clinic_list, all_documents_docx)

    # for patient in list_of_patients:
    #    print(f"For patient, {patient.CHI} the output is: {patient.output()}")
