from docx import Document
import os
from document_converter import process_documents
from collections import OrderedDict
from ClinicClass import ClinicPatient
from typing import List
from operator import attrgetter
from datetime import time, timedelta


# Pass time and CHI with list of docx documents and build the word document
def create_word_document(patient_clinic_list, file_name):
    time_sorted_clinic_list = patient_clinic_list.sort(key=lambda x: x.appt_time)
    patient_clinic_list = sorted(patient_clinic_list, key=attrgetter("appt_time"))
    doc = Document()

    for patient in patient_clinic_list:
        p = doc.add_paragraph()  # Create a new paragraph
        run = p.add_run(str(patient.title_output()))  # Add patient's summary line
        run.bold = True  # Make the appointment time bold

         # Add patient's summary as a new paragraph
        p = doc.add_paragraph(
            patient.summary
        ) 
        #Optional last seen paragraph
        #p = doc.add_paragraph(
            #patient.last_seen
        #)  
    
    #hardcoded export directory
    directory= "prepped clinics"

    doc.save(os.path.join(directory,file_name))
    return doc

if __name__ == "__main__":
    print("Running Test with doc builder")

